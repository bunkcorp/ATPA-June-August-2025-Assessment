#!/usr/bin/env python3
"""
Convert Word document to Markdown and structured JSON
"""

import json
import re
from pathlib import Path

try:
    from docx import Document
    from markdownify import markdownify as md
except ImportError:
    print("Required packages not installed. Please run:")
    print("pip install python-docx markdownify")
    exit(1)

def convert_doc_to_markdown(doc_path, markdown_path):
    """Convert Word doc to Markdown format"""
    doc = Document(doc_path)
    full_text = []
    
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            # Check if it's a heading based on style
            if para.style.name.startswith('Heading'):
                level = int(para.style.name.split()[-1]) if para.style.name.split()[-1].isdigit() else 1
                full_text.append('#' * level + ' ' + text)
            else:
                full_text.append(text)
    
    combined = "\n\n".join(full_text)
    
    with open(markdown_path, "w", encoding="utf-8") as f:
        f.write(combined)
    
    return combined

def doc_to_structured_json(doc_path):
    """Convert Word doc to structured JSON by parsing sections"""
    doc = Document(doc_path)
    sections = {}
    current_section = None
    current_content = []
    
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
            
        # Check for section headers (various patterns)
        header_patterns = [
            r"^(\d+(\.\d+)*)\s+(.+)",  # 1.1 Title, 2.3.4 Title
            r"^(Module\s+\d+[:\s]+.+)",  # Module 1: Title
            r"^(Section\s+\d+[:\s]+.+)",  # Section 1: Title
            r"^([A-Z][^a-z]*[A-Z])$",  # ALL CAPS titles
        ]
        
        is_header = False
        for pattern in header_patterns:
            match = re.match(pattern, text, re.IGNORECASE)
            if match:
                # Save previous section
                if current_section and current_content:
                    sections[current_section] = "\n".join(current_content)
                
                # Start new section
                current_section = text
                current_content = []
                is_header = True
                break
        
        if not is_header:
            current_content.append(text)
    
    # Save final section
    if current_section and current_content:
        sections[current_section] = "\n".join(current_content)
    
    return sections

def extract_meta_information(doc_path):
    """Extract titles, descriptions, and key information"""
    doc = Document(doc_path)
    meta_info = {
        "title": "",
        "sections": [],
        "learning_objectives": [],
        "key_terms": [],
        "full_text": ""
    }
    
    all_text = []
    current_section = None
    
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
            
        all_text.append(text)
        
        # Try to identify document title (first heading or bold text)
        if not meta_info["title"] and (para.style.name.startswith('Heading') or 
                                      any(run.bold for run in para.runs)):
            meta_info["title"] = text
        
        # Look for learning objectives
        if re.search(r"learning\s+objective|objective", text, re.IGNORECASE):
            meta_info["learning_objectives"].append(text)
        
        # Look for section headers
        if re.match(r"^\d+(\.\d+)*\s+", text) or para.style.name.startswith('Heading'):
            meta_info["sections"].append(text)
    
    meta_info["full_text"] = "\n".join(all_text)
    return meta_info

def main():
    # File paths
    doc_path = "/Users/kevinwoods/Desktop/ActuarialExams/ATPA/Module_1_Data_and_Model_Ethics/ATPA_Module_1_document.doc"
    base_path = Path(doc_path).parent
    
    output_files = {
        "markdown": base_path / "ATPA_Module_1_content.md",
        "json_structured": base_path / "ATPA_Module_1_structured.json",
        "json_meta": base_path / "ATPA_Module_1_meta.json"
    }
    
    print(f"Converting: {doc_path}")
    
    try:
        # Convert to Markdown
        print("Converting to Markdown...")
        markdown_content = convert_doc_to_markdown(doc_path, output_files["markdown"])
        print(f"✅ Markdown saved to: {output_files['markdown']}")
        
        # Convert to structured JSON
        print("Converting to structured JSON...")
        structured_data = doc_to_structured_json(doc_path)
        with open(output_files["json_structured"], "w", encoding="utf-8") as f:
            json.dump(structured_data, f, indent=2, ensure_ascii=False)
        print(f"✅ Structured JSON saved to: {output_files['json_structured']}")
        
        # Extract meta information
        print("Extracting meta information...")
        meta_data = extract_meta_information(doc_path)
        with open(output_files["json_meta"], "w", encoding="utf-8") as f:
            json.dump(meta_data, f, indent=2, ensure_ascii=False)
        print(f"✅ Meta JSON saved to: {output_files['json_meta']}")
        
        # Print summary
        print(f"\n📊 Summary:")
        print(f"   Title: {meta_data['title']}")
        print(f"   Sections found: {len(structured_data)}")
        print(f"   Learning objectives: {len(meta_data['learning_objectives'])}")
        print(f"   Total content length: {len(meta_data['full_text'])} characters")
        
        # Show first few sections
        if structured_data:
            print(f"\n📋 First few sections:")
            for i, section in enumerate(list(structured_data.keys())[:5]):
                print(f"   {i+1}. {section}")
        
    except Exception as e:
        print(f"❌ Error: {e}")
        print("Make sure the file exists and you have the required packages installed:")
        print("pip install python-docx markdownify")

if __name__ == "__main__":
    main()