#!/usr/bin/env python3
"""
Enhanced Word Document Manager
Combines ATPA content generation with Office-Word-MCP-Server capabilities
"""

import os
import logging
from typing import Dict, List, Any, Optional
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.shared import OxmlElement, qn

# Office-Word-MCP-Server is designed to run as a standalone MCP server
# We'll use basic functionality and provide integration points
OFFICE_WORD_AVAILABLE = False
logging.info("Using enhanced basic functionality with Office-Word-MCP-Server integration points")

from word_content_generator import WordContentGenerator

class EnhancedWordManager:
    """
    Enhanced Word document manager that combines:
    - ATPA content generation
    - Professional Word document manipulation
    - Advanced formatting and styling
    """
    
    def __init__(self):
        self.content_generator = WordContentGenerator()
        self.word_server = None
        
        if OFFICE_WORD_AVAILABLE:
            try:
                self.word_server = WordDocumentServer()
                logging.info("Office-Word-MCP-Server initialized successfully")
            except Exception as e:
                logging.error(f"Failed to initialize Office-Word-MCP-Server: {e}")
                self.word_server = None
    
    def create_professional_atpa_document(self, filename: str = "ATPA_Professional_Submission.docx") -> str:
        """Create a professionally formatted ATPA document"""
        if self.word_server:
            return self._create_with_office_word_server(filename)
        else:
            return self._create_with_basic_functionality(filename)
    
    def _create_with_office_word_server(self, filename: str) -> str:
        """Create document using Office-Word-MCP-Server capabilities"""
        try:
            # Create new document with metadata
            self.word_server.create_document(
                filename=filename,
                title="ATPA Assessment - June to August 2025",
                author="ATPA Candidate"
            )
            
            # Add title page
            self.word_server.add_heading(filename, "ATPA ASSESSMENT", 0)
            self.word_server.add_paragraph(filename, "June to August 2025 Assessment")
            self.word_server.add_paragraph(filename, "ATPA Candidate")
            self.word_server.add_page_break(filename)
            
            # Add instructions
            self.word_server.add_heading(filename, "Instructions", 1)
            instructions = """You must enter your answers to each assessment question in the sections noted below, and must not change any information contained within the first set of black brackets [] for each Task."""
            self.word_server.add_paragraph(filename, instructions)
            
            # Add table instructions
            self._add_table_instructions_with_server(filename)
            
            # Add all tasks with professional formatting
            for task_num in range(1, 7):
                self._add_task_with_server(filename, task_num)
            
            return filename
            
        except Exception as e:
            logging.error(f"Error creating document with Office-Word-MCP-Server: {e}")
            return self._create_with_basic_functionality(filename)
    
    def _create_with_basic_functionality(self, filename: str) -> str:
        """Create document using basic python-docx functionality"""
        doc = Document()
        
        # Setup document styles
        self._setup_document_styles(doc)
        
        # Add title
        title = doc.add_heading('ATPA ASSESSMENT', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add subtitle
        subtitle = doc.add_paragraph('June to August 2025 Assessment')
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add instructions
        doc.add_heading('Instructions', 1)
        instructions = doc.add_paragraph("You must enter your answers to each assessment question in the sections noted below, and must not change any information contained within the first set of black brackets [] for each Task.")
        
        # Add table instructions
        self._add_table_instructions_basic(doc)
        
        # Add all tasks
        for task_num in range(1, 7):
            self._add_task_basic(doc, task_num)
        
        # Save document
        doc.save(filename)
        return filename
    
    def _setup_document_styles(self, doc: Document):
        """Setup professional document styles"""
        # Set default font
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Calibri'
        font.size = Pt(11)
        
        # Create heading styles
        for i in range(1, 4):
            heading_style = doc.styles[f'Heading {i}']
            font = heading_style.font
            font.name = 'Calibri'
            font.size = Pt(16 - i)
            font.bold = True
    
    def _add_table_instructions_with_server(self, filename: str):
        """Add table instructions using Office-Word-MCP-Server"""
        self.word_server.add_heading(filename, "Important Notes about Inserting Tables", 2)
        
        # Add instructions with proper formatting
        instructions = [
            "If you need to insert an Excel table into your document, then you must follow the instructions below to copy and paste the table from Excel into your Word document.",
            "",
            "Do:",
            "- Ensure your table fits in the page margins by:",
            "  - Selecting the table",
            "  - Right-clicking on the table to view the pop-up menu",
            "  - Hover your mouse over \"Auto Fit\"",
            "  - Click \"AutoFit to Window\".",
            "",
            "Do not:",
            "- Paste as Picture",
            "- Insert Object",
            "- Use Word's \"Insert, Table, Excel Spreadsheet\" function",
            "",
            "Pasting a table as a picture or as an object will result in an automatic disqualification of your submission. Graphs pasted as pictures are not affected by this requirement."
        ]
        
        for instruction in instructions:
            if instruction.startswith("-"):
                # Format as bullet point
                self.word_server.add_paragraph(filename, instruction, style="List Bullet")
            elif instruction.startswith("Do:") or instruction.startswith("Do not:"):
                # Format as bold
                self.word_server.format_text(filename, 0, 0, len(instruction), bold=True)
                self.word_server.add_paragraph(filename, instruction)
            elif instruction == "":
                self.word_server.add_paragraph(filename, "")
            else:
                self.word_server.add_paragraph(filename, instruction)
    
    def _add_table_instructions_basic(self, doc: Document):
        """Add table instructions using basic functionality"""
        doc.add_heading("Important Notes about Inserting Tables", 2)
        
        instructions = [
            "If you need to insert an Excel table into your document, then you must follow the instructions below to copy and paste the table from Excel into your Word document.",
            "",
            "Do:",
            "- Ensure your table fits in the page margins by:",
            "  - Selecting the table",
            "  - Right-clicking on the table to view the pop-up menu",
            "  - Hover your mouse over \"Auto Fit\"",
            "  - Click \"AutoFit to Window\".",
            "",
            "Do not:",
            "- Paste as Picture",
            "- Insert Object",
            "- Use Word's \"Insert, Table, Excel Spreadsheet\" function",
            "",
            "Pasting a table as a picture or as an object will result in an automatic disqualification of your submission. Graphs pasted as pictures are not affected by this requirement."
        ]
        
        for instruction in instructions:
            if instruction.startswith("-"):
                doc.add_paragraph(instruction, style="List Bullet")
            elif instruction.startswith("Do:") or instruction.startswith("Do not:"):
                p = doc.add_paragraph(instruction)
                p.runs[0].bold = True
            elif instruction == "":
                doc.add_paragraph()
            else:
                doc.add_paragraph(instruction)
    
    def _add_task_with_server(self, filename: str, task_num: int):
        """Add task content using Office-Word-MCP-Server"""
        task_headers = {
            1: "Task 1",
            2: "Task 2", 
            3: "Task 3",
            4: "Task 4",
            5: "Task 5",
            6: "Task 6"
        }
        
        # Add task header
        self.word_server.add_heading(filename, f"[{task_headers[task_num]}]", 1)
        
        # Add atpaksat identifier
        self.word_server.add_paragraph(filename, "atpaksat")
        
        # Get task content
        content = self._get_task_content(task_num)
        
        # Add content with proper formatting
        lines = content.split('\n')
        for line in lines:
            if line.strip():
                if line.startswith('DATA PREPARATION') or line.startswith('PRIVACY') or line.startswith('GENERALIZED') or line.startswith('RANDOM') or line.startswith('BAYESIAN') or line.startswith('EXECUTIVE'):
                    # Main section headers
                    self.word_server.add_heading(filename, line, 2)
                elif line.startswith('a)') or line.startswith('b)') or line.startswith('c)') or line.startswith('d)') or line.startswith('e)'):
                    # Subsection headers
                    self.word_server.add_heading(filename, line, 3)
                elif line.startswith('-'):
                    # Bullet points
                    self.word_server.add_paragraph(filename, line, style="List Bullet")
                else:
                    # Regular paragraphs
                    self.word_server.add_paragraph(filename, line)
    
    def _add_task_basic(self, doc: Document, task_num: int):
        """Add task content using basic functionality"""
        task_headers = {
            1: "Task 1",
            2: "Task 2", 
            3: "Task 3",
            4: "Task 4",
            5: "Task 5",
            6: "Task 6"
        }
        
        # Add task header
        doc.add_heading(f"[{task_headers[task_num]}]", 1)
        
        # Add atpaksat identifier
        doc.add_paragraph("atpaksat")
        
        # Get task content
        content = self._get_task_content(task_num)
        
        # Add content with proper formatting
        lines = content.split('\n')
        for line in lines:
            if line.strip():
                if line.startswith('DATA PREPARATION') or line.startswith('PRIVACY') or line.startswith('GENERALIZED') or line.startswith('RANDOM') or line.startswith('BAYESIAN') or line.startswith('EXECUTIVE'):
                    # Main section headers
                    doc.add_heading(line, 2)
                elif line.startswith('a)') or line.startswith('b)') or line.startswith('c)') or line.startswith('d)') or line.startswith('e)'):
                    # Subsection headers
                    doc.add_heading(line, 3)
                elif line.startswith('-'):
                    # Bullet points
                    doc.add_paragraph(line, style="List Bullet")
                else:
                    # Regular paragraphs
                    doc.add_paragraph(line)
    
    def _get_task_content(self, task_num: int) -> str:
        """Get content for a specific task"""
        task_content_map = {
            1: self.content_generator.generate_task1_content(),
            2: self.content_generator.generate_task2_content(),
            3: self.content_generator.generate_task3_content(),
            4: self.content_generator.generate_task4_content(),
            5: self.content_generator.generate_task5_content(),
            6: self.content_generator.generate_task6_content()
        }
        return task_content_map.get(task_num, "")
    
    def update_existing_template(self, template_path: str) -> str:
        """Update existing template with enhanced formatting"""
        if self.word_server:
            return self._update_with_office_word_server(template_path)
        else:
            return self._update_with_basic_functionality(template_path)
    
    def _update_with_office_word_server(self, template_path: str) -> str:
        """Update template using Office-Word-MCP-Server - preserving original structure"""
        try:
            # For now, fall back to basic functionality to ensure we preserve original content
            return self._update_with_basic_functionality(template_path)
            
        except Exception as e:
            logging.error(f"Error updating with Office-Word-MCP-Server: {e}")
            return self._update_with_basic_functionality(template_path)
    
    def _update_with_basic_functionality(self, template_path: str) -> str:
        """Update template using basic functionality"""
        return self.content_generator.update_template_with_content(template_path)
    
    def add_table_to_document(self, filename: str, table_data: List[List[str]], table_title: str = "") -> bool:
        """Add a professionally formatted table to the document"""
        if self.word_server:
            return self._add_table_with_server(filename, table_data, table_title)
        else:
            return self._add_table_basic(filename, table_data, table_title)
    
    def _add_table_with_server(self, filename: str, table_data: List[List[str]], table_title: str) -> bool:
        """Add table using Office-Word-MCP-Server"""
        try:
            if table_title:
                self.word_server.add_heading(filename, table_title, 3)
            
            rows = len(table_data)
            cols = len(table_data[0]) if table_data else 0
            
            self.word_server.add_table(filename, rows, cols, table_data)
            
            # Format table professionally
            self.word_server.format_table(filename, 0, has_header_row=True, border_style="single", shading="light_gray")
            
            return True
            
        except Exception as e:
            logging.error(f"Error adding table with server: {e}")
            return False
    
    def _add_table_basic(self, filename: str, table_data: List[List[str]], table_title: str) -> bool:
        """Add table using basic functionality"""
        try:
            doc = Document(filename)
            
            if table_title:
                doc.add_heading(table_title, 3)
            
            if table_data:
                table = doc.add_table(rows=len(table_data), cols=len(table_data[0]))
                table.style = 'Table Grid'
                
                for i, row_data in enumerate(table_data):
                    for j, cell_data in enumerate(row_data):
                        table.cell(i, j).text = str(cell_data)
            
            doc.save(filename)
            return True
            
        except Exception as e:
            logging.error(f"Error adding table: {e}")
            return False
    
    def format_text_in_document(self, filename: str, paragraph_index: int, start_pos: int, end_pos: int, 
                               bold: bool = None, italic: bool = None, color: str = None) -> bool:
        """Format specific text in the document"""
        if self.word_server:
            return self._format_text_with_server(filename, paragraph_index, start_pos, end_pos, bold, italic, color)
        else:
            return self._format_text_basic(filename, paragraph_index, start_pos, end_pos, bold, italic, color)
    
    def _format_text_with_server(self, filename: str, paragraph_index: int, start_pos: int, end_pos: int,
                                bold: bool = None, italic: bool = None, color: str = None) -> bool:
        """Format text using Office-Word-MCP-Server"""
        try:
            self.word_server.format_text(filename, paragraph_index, start_pos, end_pos, 
                                       bold=bold, italic=italic, color=color)
            return True
        except Exception as e:
            logging.error(f"Error formatting text with server: {e}")
            return False
    
    def _format_text_basic(self, filename: str, paragraph_index: int, start_pos: int, end_pos: int,
                          bold: bool = None, italic: bool = None, color: str = None) -> bool:
        """Format text using basic functionality"""
        try:
            doc = Document(filename)
            
            if paragraph_index < len(doc.paragraphs):
                paragraph = doc.paragraphs[paragraph_index]
                
                if start_pos < len(paragraph.text) and end_pos <= len(paragraph.text):
                    # This is a simplified version - full implementation would be more complex
                    run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
                    
                    if bold is not None:
                        run.bold = bold
                    if italic is not None:
                        run.italic = italic
                    if color:
                        run.font.color.rgb = color
            
            doc.save(filename)
            return True
            
        except Exception as e:
            logging.error(f"Error formatting text: {e}")
            return False
    
    def convert_to_pdf(self, filename: str) -> str:
        """Convert Word document to PDF"""
        if self.word_server:
            try:
                pdf_filename = filename.replace('.docx', '.pdf')
                self.word_server.convert_to_pdf(filename, pdf_filename)
                return pdf_filename
            except Exception as e:
                logging.error(f"Error converting to PDF: {e}")
                return filename
        else:
            logging.warning("PDF conversion not available without Office-Word-MCP-Server")
            return filename

def create_enhanced_atpa_document() -> str:
    """Create an enhanced ATPA document with professional formatting"""
    manager = EnhancedWordManager()
    return manager.create_professional_atpa_document()

if __name__ == "__main__":
    # Test the enhanced word manager
    manager = EnhancedWordManager()
    filename = manager.create_professional_atpa_document()
    print(f"Created enhanced ATPA document: {filename}") 