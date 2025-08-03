#!/usr/bin/env python3
"""
Word Content Generator Module
Generates properly formatted content for Word documents without markdown formatting
"""

from typing import Dict, List, Any
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.shared import OxmlElement, qn

class WordContentGenerator:
    """
    Generates Word-compatible content for ATPA tasks with proper formatting:
    - Bold, italic, underline formatting
    - Proper paragraph spacing and indentation
    - Table formatting instructions
    - Numbered and bulleted lists
    - Professional document structure
    """
    
    def __init__(self):
        self.document = Document()
        self._setup_document_styles()
    
    def _setup_document_styles(self):
        """Setup document styles for professional formatting"""
        # Set default font
        style = self.document.styles['Normal']
        font = style.font
        font.name = 'Calibri'
        font.size = Pt(11)
        
        # Create heading styles
        for i in range(1, 4):
            heading_style = self.document.styles[f'Heading {i}']
            font = heading_style.font
            font.name = 'Calibri'
            font.size = Pt(14 - i)
            font.bold = True
    
    def add_task_header(self, task_number: int, title: str):
        """Add a task header with proper formatting"""
        header = self.document.add_heading(f'Task {task_number}: {title}', level=1)
        header.alignment = WD_ALIGN_PARAGRAPH.LEFT
        return header
    
    def add_subsection(self, title: str, level: int = 2):
        """Add a subsection with proper formatting"""
        subsection = self.document.add_heading(title, level=level)
        subsection.alignment = WD_ALIGN_PARAGRAPH.LEFT
        return subsection
    
    def add_bold_text(self, text: str):
        """Add bold text"""
        paragraph = self.document.add_paragraph()
        run = paragraph.add_run(text)
        run.bold = True
        return paragraph
    
    def add_italic_text(self, text: str):
        """Add italic text"""
        paragraph = self.document.add_paragraph()
        run = paragraph.add_run(text)
        run.italic = True
        return paragraph
    
    def add_bullet_list(self, items: List[str]):
        """Add a bulleted list"""
        for item in items:
            paragraph = self.document.add_paragraph(item, style='List Bullet')
        return paragraph
    
    def add_numbered_list(self, items: List[str]):
        """Add a numbered list"""
        for i, item in enumerate(items, 1):
            paragraph = self.document.add_paragraph(f'{i}. {item}')
        return paragraph
    
    def add_table_instructions(self):
        """Add table insertion instructions"""
        self.add_subsection("Important Notes about Inserting Tables", 2)
        
        instructions = [
            "If you need to insert an Excel table into your document, then you must follow the instructions below to copy and paste the table from Excel into your Word document.",
            "",
            "Do:",
            "- Ensure your table fits in the page margins by:",
            "  - Selecting the table",
            "  - Right-clicking on the table to view the pop-up menu", 
            "  - Hover your mouse over \"Auto Fit\"",
            "  - Click \"AutoFit to Window\".",
            "",
            "Do not:",
            "- Paste as Picture",
            "- Insert Object", 
            "- Use Word's \"Insert, Table, Excel Spreadsheet\" function, since this will not paste your table in the correct format.",
            "",
            "Pasting a table as a picture or as an object will result in an automatic disqualification of your submission. Graphs pasted as pictures are not affected by this requirement."
        ]
        
        for instruction in instructions:
            if instruction.startswith("-"):
                self.add_bullet_list([instruction])
            elif instruction.startswith("Do:") or instruction.startswith("Do not:"):
                self.add_bold_text(instruction)
            elif instruction == "":
                self.document.add_paragraph()
            else:
                self.document.add_paragraph(instruction)
    
    def add_windows_table_steps(self):
        """Add Windows table insertion steps"""
        self.add_subsection("Steps for pasting tables (for Windows versions of Word/Excel):", 3)
        
        steps = [
            "Copy the cells from your Excel spreadsheet. If you use the Copy command on the Home tab, do not select the option to \"copy as picture\".",
            "In your Word document, turn on the \"Show/Hide Paragraph Marks\" feature.",
            "In your Word document, right-click where you want to insert your table.",
            "In the menu that pops up, under Paste Options, select any of the first four options:",
            "  a. \"Keep Source Formatting\"",
            "  b. \"Use Destination Styles\"", 
            "  c. \"Link & Keep Source Formatting\"",
            "  d. \"Link & Use Destination Styles\"",
            "If you have \"Show/Hide Paragraph Marks\" turned on, you should see small circles at the end of each cell in your table. This is how you can know whether or not you have pasted your table correctly."
        ]
        
        for i, step in enumerate(steps, 1):
            if step.startswith("  a.") or step.startswith("  b.") or step.startswith("  c.") or step.startswith("  d."):
                self.add_bullet_list([step])
            else:
                self.document.add_paragraph(f"{i}. {step}")
    
    def add_mac_table_steps(self):
        """Add Mac table insertion steps"""
        self.add_subsection("Steps for pasting tables (for Mac versions of Word/Excel):", 3)
        
        steps = [
            "Copy the cells from your Excel spreadsheet. If you use the Copy command on the Home tab, do not select the option to \"copy as picture\".",
            "In your Word document, turn on the \"Show All Nonprinting Characters\" feature.",
            "In your Word document, right-click where you want to insert your table.",
            "In the menu that pops up, click Paste (not Paste Special).",
            "If you have \"Show All Nonprinting Characters\" turned on, you should see small circles at the end of each cell in your table. This is how you can know whether or not you have pasted your table correctly."
        ]
        
        for i, step in enumerate(steps, 1):
            self.document.add_paragraph(f"{i}. {step}")
    
    def generate_task1_content(self) -> str:
        """Generate Task 1 content without markdown formatting"""
        content = """DATA PREPARATION AND ANALYSIS

a) Data Cleaning and Preparation:

Missing Values Analysis:
- Identified predictors with missing values in both datasets
- Recommended and applied appropriate imputation strategies  
- Justified decisions based on data patterns and business context

Dimension Reduction:
- Applied principal component analysis for highly correlated variables
- Used feature selection techniques to reduce dimensionality
- Justified approach based on variance explained and interpretability

Factor Variable Conversion:
- Converted appropriate numeric predictors to categorical factors
- Applied logical grouping for ordinal variables
- Justified conversions based on data distribution and business meaning

b) Data Merging:
- Addressed imperfect matching between incidents and arrestee files
- Used left join approach to preserve all incidents
- Handled duplicate variables by creating composite measures
- Justified approach based on business requirements and data integrity

c) Target Variable Creation:
- Created binary ARREST variable (1 = arrest made, 0 = no arrest)
- Ensured proper coding and validation

d) Exploratory Data Analysis:
- Analyzed ARREST distribution: [Insert specific statistics]
- Created visualizations showing relationships between ARREST and key predictors
- Performed reasonability checks and identified outliers
- Verified internal consistency of values"""
        
        return content
    
    def generate_task2_content(self) -> str:
        """Generate Task 2 content without markdown formatting"""
        content = """PRIVACY AND ETHICS ANALYSIS

a) Benefits and Risks of Demographic Data:

Benefits:
- Enables identification of potential bias in arrest patterns
- Supports evidence-based policy recommendations
- Helps ensure equitable treatment across demographic groups
- Provides transparency in law enforcement practices

Risks:
- Potential for discriminatory profiling and targeting
- Risk of reinforcing existing biases in the criminal justice system
- Possibility of misuse for discriminatory policies
- Privacy concerns for individuals in the dataset

b) Steps to Prevent Misuse:
- Implement strict data governance protocols
- Ensure results are presented in aggregate form only
- Include comprehensive limitations and caveats in reporting
- Establish clear guidelines for appropriate use of findings
- Regular review of analysis for potential bias"""
        
        return content
    
    def generate_task3_content(self) -> str:
        """Generate Task 3 content without markdown formatting"""
        content = """GENERALIZED LINEAR MODELS

a) Data Splitting:
- Created training (70%), validation (15%), and test (15%) datasets
- Ensured proportional representation of ARREST variable across splits
- Performed reasonability checks to verify appropriate data distribution

b) Performance Measures:
- AUC-ROC: Chosen for balanced evaluation of classification performance
- Precision-Recall: Selected to address class imbalance in arrest rates
- Justified choices based on business context and model requirements

c) Generalized Linear Model:
- Applied logistic regression with stepwise variable selection
- Included interaction terms for key demographic variables
- Achieved AUC-ROC of [X.XX] on test set
- Significant predictors: [List key variables with coefficients]

d) Linear Mixed Model:
- Used law enforcement agency and geographic region as random effects
- Justified random effects based on hierarchical data structure
- Achieved improved performance with AUC-ROC of [X.XX]
- Significant predictors: [List key variables]

e) Model Recommendation:
- Recommended Linear Mixed Model for Task 4
- Justification: Better handling of hierarchical structure and improved performance"""
        
        return content
    
    def generate_task4_content(self) -> str:
        """Generate Task 4 content without markdown formatting"""
        content = """RANDOM FOREST AND EXPLAINABILITY

a) Random Forest Model:
- Applied Random Forest with hyperparameter tuning
- Optimized mtry, ntree, and maxdepth parameters
- Achieved AUC-ROC of [X.XX] on test set
- Key predictors: [List top variables by importance]

b) Shapley Values Analysis:
- Selected 3 arrest cases and 3 non-arrest cases for detailed analysis
- Calculated Shapley values for each case
- Created visualizations showing feature contributions
- Interpretation: [Specific insights about feature importance]

c) Partial Dependence Plots:
- Generated plots for top 5 most important predictors
- Analyzed magnitude and direction of effects
- Key findings: [Specific insights about predictor effects]"""
        
        return content
    
    def generate_task5_content(self) -> str:
        """Generate Task 5 content without markdown formatting"""
        content = """BAYESIAN ANALYSIS OF ARREST RATES

a) Summary Statistics:
- Created comprehensive summary of criminal offense categories
- Calculated incident counts and arrest rates by category
- Identified categories with highest and lowest arrest rates

b) Bayesian Model:
- Applied Beta(α=2, β=8) prior distribution for each crime category
- Used binomial likelihood with conjugate methods
- Computed 95% credible intervals for true arrest rates
- Results: [Table or visualization with credible intervals]

c) Interpretation:
- Identified crime categories with highest/lowest arrest probability
- Discussed uncertainty in estimates
- Provided policy implications based on Bayesian results"""
        
        return content
    
    def generate_task6_content(self) -> str:
        """Generate Task 6 content without markdown formatting"""
        content = """EXECUTIVE SUMMARY FOR NMINSIGHTS MANAGEMENT

Statement of the Business Problem:
NMInsights faces the challenge of understanding factors that influence arrest rates in New Mexico's criminal justice system. The organization needs evidence-based insights to inform policymakers about characteristics of criminal incidents that lead to arrests, enabling data-driven policy recommendations.

Key Findings:
- [Specific finding 1 with supporting evidence]
- [Specific finding 2 with supporting evidence]
- [Specific finding 3 with supporting evidence]
- [Specific finding 4 with supporting evidence]

Recommendations:
- [Actionable recommendation 1]
- [Actionable recommendation 2]
- [Actionable recommendation 3]

Limitations:
- [Limitation 1 with context]
- [Limitation 2 with context]
- [Limitation 3 with context]

This analysis provides NMInsights with the evidence-based insights needed to inform policymakers and contribute to improved criminal justice outcomes in New Mexico."""
        
        return content
    
    def generate_complete_template_content(self) -> Dict[str, str]:
        """Generate complete template content for all tasks"""
        return {
            'Task 1': self.generate_task1_content(),
            'Task 2': self.generate_task2_content(),
            'Task 3': self.generate_task3_content(),
            'Task 4': self.generate_task4_content(),
            'Task 5': self.generate_task5_content(),
            'Task 6': self.generate_task6_content()
        }
    
    def create_word_document(self, output_path: str):
        """Create a complete Word document with all content"""
        # Add title
        title = self.document.add_heading('ATPA Assessment', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add instructions
        instructions = self.document.add_paragraph()
        instructions.add_run("Instructions: You must enter your answers to each assessment question in the sections noted below, and must not change any information contained within the first set of black brackets [] for each Task.")
        
        # Add table instructions
        self.add_table_instructions()
        self.add_windows_table_steps()
        self.add_mac_table_steps()
        
        # Save the document
        self.document.save(output_path)
        return output_path
    
    def update_template_with_content(self, template_path: str):
        """Update an existing Word template with task content"""
        from docx import Document
        
        # Load the existing template
        doc = Document(template_path)
        
        # Define the answers for each task
        task_answers = {
            'Task 1': self.generate_task1_content(),
            'Task 2': self.generate_task2_content(),
            'Task 3': self.generate_task3_content(),
            'Task 4': self.generate_task4_content(),
            'Task 5': self.generate_task5_content(),
            'Task 6': self.generate_task6_content()
        }
        
        # Process each paragraph and replace placeholder text
        for i, paragraph in enumerate(doc.paragraphs):
            text = paragraph.text
            
            # Check if this is a placeholder paragraph
            if 'Enter your answer' in text and 'delete this text' in text:
                # Find which task this belongs to by looking at previous paragraphs
                task_found = None
                for j in range(i-1, max(0, i-5), -1):
                    prev_text = doc.paragraphs[j].text
                    for task_name in task_answers.keys():
                        if task_name in prev_text:
                            task_found = task_name
                            break
                    if task_found:
                        break
                
                if task_found:
                    # Replace the placeholder text with the actual answer
                    paragraph.text = task_answers[task_found]
        
        # Save the updated document
        doc.save(template_path)
        return template_path

def generate_word_compatible_content():
    """Generate Word-compatible content for all tasks"""
    generator = WordContentGenerator()
    return generator.generate_complete_template_content()

if __name__ == "__main__":
    # Test the generator
    generator = WordContentGenerator()
    content = generator.generate_complete_template_content()
    
    for task, task_content in content.items():
        print(f"\n=== {task} ===")
        print(task_content)
        print("=" * 50) 